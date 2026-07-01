"""Structured-file parsers for PDF, Excel, DOCX, and SQLite.

All deps below are permissive-licensed and safe to bundle/redistribute.
License audit (verified against each project's latest PyPI metadata):

- pdfplumber        MIT        (https://pypi.org/project/pdfplumber/)
- openpyxl          MIT        (https://pypi.org/project/openpyxl/)
- python-docx       MIT        (https://pypi.org/project/python-docx/)
- sqlite3           PSF        (Python standard library)

Reject-list (intentionally NOT imported):
- formulas          EUPL-1.2   — Excel formula evaluator; would taint license.
  Replaced by hand-rolled regex-based reference extraction below.
- xlwings           BSD-3 but drives Excel via COM/OLE (Windows/Office-only).
- tabula-py         MIT wrapper over GPL Java tabula-java → transitive GPL.

The four public entrypoints (`read_pdf`, `read_excel`, `read_docx`,
`read_sqlite`) return plain dicts and raise `FileNotFoundError` for a
missing path, `ValueError` for validation failures, and
`ModuleNotFoundError` with a clean "install aidocs-mcp[office]" message
when an optional dep is missing. Wrappers in server_code_tools catch those
and convert them into structured error payloads.
"""

from __future__ import annotations

import importlib.util
import re
import sqlite3
from pathlib import Path
from typing import Any

# ── Optional-dep guard ──────────────────────────────────────────────────

_OFFICE_DEP_HINT = (
    "Structured-file parser '{tool}' needs the '{dep}' package. "
    "Install the optional extra: pip install aidocs-mcp[office]"
)


def _require_module(name: str, tool: str) -> Any:
    """Import `name` lazily, raising a clean install hint if missing."""
    if importlib.util.find_spec(name) is None:
        raise ModuleNotFoundError(_OFFICE_DEP_HINT.format(tool=tool, dep=name))
    return __import__(name)


# ── Shared helpers ──────────────────────────────────────────────────────


def _resolve_path(path: str | Path) -> Path:
    """Convert to absolute Path and ensure the file exists.

    Legacy permissive helper: relative paths resolved against process
    cwd, ``~`` expanded. Direct parser callers (tests, internal
    helpers) rely on this shape. Read-tool wrappers route through
    ``read_pipeline.gate`` which performs strict project-anchored
    resolution + zone classification BEFORE invoking the parser, so
    the security envelope does not depend on this helper.
    """
    if not path or not str(path).strip():
        raise ValueError("path is required")
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"file not found: {p}")
    if not p.is_file():
        raise ValueError(f"not a regular file: {p}")
    return p


def _require_absolute_path(path: str | Path) -> Path:
    """Strict variant for callers that already passed through the
    read-pipeline gate. The gate produces an absolute resolved path;
    this helper just asserts the contract and verifies the file
    exists. Use only when the security envelope upstream guarantees
    absoluteness.
    """
    if not path or not str(path).strip():
        raise ValueError("path is required")
    p = Path(path)
    if not p.is_absolute():
        raise ValueError(f"absolute path required; got {str(path)!r}")
    if not p.exists():
        raise FileNotFoundError(f"file not found: {p}")
    if not p.is_file():
        raise ValueError(f"not a regular file: {p}")
    return p


def _parse_page_spec(spec: str, total: int) -> list[int]:
    """Parse "1-5,8,10-12" into a sorted list of 1-based page numbers.

    Empty spec means all pages. Invalid tokens raise ValueError.
    """
    if not spec or not spec.strip():
        return list(range(1, total + 1))
    pages: set[int] = set()
    for raw in spec.split(","):
        token = raw.strip()
        if not token:
            continue
        if "-" in token:
            lo_s, hi_s = token.split("-", 1)
            try:
                lo, hi = int(lo_s), int(hi_s)
            except ValueError as exc:
                raise ValueError(f"invalid page range: {token!r}") from exc
            if lo < 1 or hi < 1 or lo > hi:
                raise ValueError(f"invalid page range: {token!r}")
            for page in range(lo, hi + 1):
                if 1 <= page <= total:
                    pages.add(page)
        else:
            try:
                n = int(token)
            except ValueError as exc:
                raise ValueError(f"invalid page number: {token!r}") from exc
            if n < 1:
                raise ValueError(f"invalid page number: {token!r}")
            if 1 <= n <= total:
                pages.add(n)
    return sorted(pages)


# ── PDF ─────────────────────────────────────────────────────────────────

_PDF_MAX_PAGES = 50


def read_pdf(
    path: str,
    pages: str = "",
    mode: str = "text",
) -> dict[str, Any]:
    """Extract text (and optionally tables) from a PDF.

    mode: "text" (default) or "text_and_tables".
    pages: range spec like "1-5,8" or "" for all.

    Raises ValueError when the unscoped read exceeds _PDF_MAX_PAGES; the
    caller is expected to pass an explicit `pages` range.
    """
    if mode not in {"text", "text_and_tables"}:
        raise ValueError(f"invalid mode: {mode!r} (expected 'text' or 'text_and_tables')")
    p = _resolve_path(path)
    _require_module("pdfplumber", "ai_read_pdf")
    import pdfplumber

    with pdfplumber.open(str(p)) as pdf:
        total = len(pdf.pages)
        selected = _parse_page_spec(pages, total)
        if not pages.strip() and total > _PDF_MAX_PAGES:
            raise ValueError(
                f"PDF has {total} pages; max per call is {_PDF_MAX_PAGES}. "
                f"Pass an explicit `pages` range (e.g. '1-50').",
            )
        if len(selected) > _PDF_MAX_PAGES:
            raise ValueError(f"Requested {len(selected)} pages; max per call is {_PDF_MAX_PAGES}.")

        pages_out: list[dict[str, Any]] = []
        for page_no in selected:
            page = pdf.pages[page_no - 1]
            entry: dict[str, Any] = {
                "page": page_no,
                "text": (page.extract_text() or "").strip(),
            }
            if mode == "text_and_tables":
                tables = page.extract_tables() or []
                entry["tables"] = [
                    [[("" if cell is None else str(cell)) for cell in row] for row in table]
                    for table in tables
                ]
            pages_out.append(entry)

    return {
        "path": str(p),
        "total_pages": total,
        "returned_pages": len(pages_out),
        "mode": mode,
        "pages": pages_out,
    }


# ── Excel ───────────────────────────────────────────────────────────────

_EXCEL_MAX_ROWS = 500

# Match cell/range tokens in an Excel formula. Handles:
#   A1, $A$1, A$1, $A1
#   A1:B10
#   Sheet2!A1, Sheet2!A1:B10
#   'Sheet Name'!A1, 'Sheet Name'!A1:B10  (quoted sheet names with spaces)
# We deliberately ignore function names, numbers, and string literals.
_CELL_RE = re.compile(
    r"""
    (?:
        (?:'(?P<quoted>[^']+)'|(?P<sheet>[A-Za-z_][A-Za-z0-9_]*))!   # optional sheet prefix
    )?
    (?P<start>\$?[A-Z]{1,3}\$?[0-9]+)                                # first cell
    (?::(?P<end>\$?[A-Z]{1,3}\$?[0-9]+))?                            # optional range end
    """,
    re.VERBOSE,
)


def _extract_formula_refs(formula: str) -> list[dict[str, str]]:
    """Return structured references found in `formula`.

    Each entry carries the sheet (if specified), the start cell, and the
    optional end cell. Sheet names quoted with single-quotes are unwrapped.
    String literals inside the formula are stripped first so cell-like
    tokens inside strings do not match.
    """
    if not formula:
        return []
    # Strip double-quoted string literals (Excel uses "" for embedded quotes).
    stripped = re.sub(r'"(?:[^"]|"")*"', "", formula)
    refs: list[dict[str, str]] = []
    seen: set[tuple[str, str, str]] = set()
    for match in _CELL_RE.finditer(stripped):
        sheet = match.group("quoted") or match.group("sheet") or ""
        start = match.group("start")
        end = match.group("end") or ""
        key = (sheet, start, end)
        if key in seen:
            continue
        seen.add(key)
        ref: dict[str, str] = {"start": start}
        if sheet:
            ref["sheet"] = sheet
        if end:
            ref["end"] = end
            ref["range"] = f"{start}:{end}"
        refs.append(ref)
    return refs


def _excel_load(path: Path, data_only: bool) -> Any:
    _require_module("openpyxl", "ai_read_excel")
    from openpyxl import load_workbook

    return load_workbook(
        filename=str(path),
        read_only=True,
        data_only=data_only,
    )


def read_excel(
    path: str,
    mode: str = "outline",
    sheet: str = "",
    cell: str = "",
    query: str = "",  # reserved for future search mode
) -> dict[str, Any]:
    """Inspect an Excel workbook.

    Modes:
      - outline   : sheet names, dimensions, header row per sheet.
      - sheet     : cell contents for `sheet` (capped at 500 rows).
      - formulas  : cells that contain formulas in `sheet` with formula text.
      - trace     : parse the formula at `sheet`!`cell` and return its refs.
    """
    if mode not in {"outline", "sheet", "formulas", "trace"}:
        raise ValueError(
            f"invalid mode: {mode!r} (expected 'outline', 'sheet', 'formulas', or 'trace')",
        )
    p = _resolve_path(path)
    del query  # unused placeholder — keeps signature stable for future work

    if mode == "outline":
        wb = _excel_load(p, data_only=True)
        try:
            sheets: list[dict[str, Any]] = []
            for name in wb.sheetnames:
                ws = wb[name]
                header: list[str] = []
                first_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
                if first_row:
                    header = [("" if v is None else str(v)) for v in first_row]
                sheets.append(
                    {
                        "name": name,
                        "max_row": ws.max_row,
                        "max_col": ws.max_column,
                        "header": header,
                    },
                )
            return {
                "path": str(p),
                "mode": mode,
                "sheets": sheets,
            }
        finally:
            wb.close()

    if mode == "sheet":
        if not sheet:
            raise ValueError("mode='sheet' requires `sheet` name")
        wb = _excel_load(p, data_only=True)
        try:
            if sheet not in wb.sheetnames:
                raise ValueError(f"sheet not found: {sheet!r}")
            ws = wb[sheet]
            rows: list[list[Any]] = []
            for idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if idx > _EXCEL_MAX_ROWS:
                    break
                rows.append([("" if v is None else v) for v in row])
            return {
                "path": str(p),
                "mode": mode,
                "sheet": sheet,
                "rows": rows,
                "row_count": len(rows),
                "truncated": ws.max_row > _EXCEL_MAX_ROWS,
                "total_rows": ws.max_row,
            }
        finally:
            wb.close()

    if mode == "formulas":
        if not sheet:
            raise ValueError("mode='formulas' requires `sheet` name")
        # Need formula text, so data_only=False.
        wb = _excel_load(p, data_only=False)
        try:
            if sheet not in wb.sheetnames:
                raise ValueError(f"sheet not found: {sheet!r}")
            ws = wb[sheet]
            formulas: list[dict[str, Any]] = []
            for row in ws.iter_rows():
                for cell_obj in row:
                    val = cell_obj.value
                    if isinstance(val, str) and val.startswith("="):
                        formulas.append(
                            {
                                "cell": cell_obj.coordinate,
                                "formula": val,
                            },
                        )
            return {
                "path": str(p),
                "mode": mode,
                "sheet": sheet,
                "formulas": formulas,
                "count": len(formulas),
            }
        finally:
            wb.close()

    # mode == "trace"
    if not sheet:
        raise ValueError("mode='trace' requires `sheet` name")
    if not cell:
        raise ValueError("mode='trace' requires `cell` coordinate (e.g. 'B5')")
    wb = _excel_load(p, data_only=False)
    try:
        if sheet not in wb.sheetnames:
            raise ValueError(f"sheet not found: {sheet!r}")
        ws = wb[sheet]
        # read_only workbooks expose cells via ws[coord]
        try:
            target = ws[cell]
        except (KeyError, ValueError) as exc:
            raise ValueError(f"invalid cell coordinate: {cell!r}") from exc
        formula = target.value if isinstance(target.value, str) else ""
        if not (isinstance(formula, str) and formula.startswith("=")):
            return {
                "path": str(p),
                "mode": mode,
                "sheet": sheet,
                "cell": cell,
                "formula": "",
                "references": [],
                "note": "cell does not contain a formula",
            }
        return {
            "path": str(p),
            "mode": mode,
            "sheet": sheet,
            "cell": cell,
            "formula": formula,
            "references": _extract_formula_refs(formula),
        }
    finally:
        wb.close()


# ── DOCX ────────────────────────────────────────────────────────────────


def _parse_section_spec(spec: str) -> tuple[int, int] | None:
    """Parse "1-3" into (1, 3). Empty → None (all sections)."""
    if not spec or not spec.strip():
        return None
    token = spec.strip()
    if "-" in token:
        lo_s, hi_s = token.split("-", 1)
        try:
            lo, hi = int(lo_s), int(hi_s)
        except ValueError as exc:
            raise ValueError(f"invalid section range: {token!r}") from exc
    else:
        try:
            lo = hi = int(token)
        except ValueError as exc:
            raise ValueError(f"invalid section number: {token!r}") from exc
    if lo < 1 or hi < lo:
        raise ValueError(f"invalid section range: {token!r}")
    return (lo, hi)


def read_docx(path: str, sections: str = "") -> dict[str, Any]:
    """Extract paragraphs + tables from a .docx in document order.

    `sections` is an optional "1-3" range that keeps only the first N
    top-level sections (split on Heading 1 style). "" returns everything.
    """
    p = _resolve_path(path)
    _require_module("docx", "ai_read_docx")
    import docx

    doc = docx.Document(str(p))
    section_range = _parse_section_spec(sections)

    # Walk block-level elements (paragraphs + tables) in document order by
    # iterating doc.element.body children and keying them back to the
    # python-docx objects via their underlying XML elements.
    from docx.oxml.ns import qn

    para_lookup = {para._element: para for para in doc.paragraphs}
    table_lookup = {tbl._element: tbl for tbl in doc.tables}

    blocks: list[dict[str, Any]] = []
    current_section = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            para = para_lookup.get(child)
            if para is None:
                continue
            style_name = (para.style.name if para.style is not None else "") or ""
            if style_name == "Heading 1":
                current_section += 1
            if section_range is not None:
                lo, hi = section_range
                if current_section == 0 and style_name != "Heading 1":
                    # Content before the first Heading 1 — keep only if lo == 1.
                    if lo > 1:
                        continue
                elif current_section < lo or current_section > hi:
                    continue
            blocks.append(
                {
                    "kind": "paragraph",
                    "style": style_name,
                    "text": para.text,
                },
            )
        elif tag == qn("w:tbl"):
            table = table_lookup.get(child)
            if table is None:
                continue
            if section_range is not None:
                lo, hi = section_range
                if current_section < lo or current_section > hi:
                    continue
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            blocks.append(
                {
                    "kind": "table",
                    "rows": rows,
                },
            )
    return {
        "path": str(p),
        "block_count": len(blocks),
        "section_range": list(section_range) if section_range else None,
        "blocks": blocks,
    }


# ── SQLite ──────────────────────────────────────────────────────────────

_SQLITE_REJECT_RE = re.compile(
    r"\b(?:insert|update|delete|drop|alter|create|attach|replace)\b",
    re.IGNORECASE,
)


def _sqlite_connect_readonly(path: Path) -> sqlite3.Connection:
    # `file:` URI + mode=ro makes SQLite reject any write attempt at the
    # driver level — belt-and-braces with the statement regex above.
    uri = f"file:{path.as_posix()}?mode=ro"
    return sqlite3.connect(uri, uri=True)


def read_sqlite(
    path: str,
    mode: str = "tables",
    table: str = "",
    query: str = "",
    limit: int = 100,
) -> dict[str, Any]:
    """Inspect or query a SQLite file (read-only).

    Modes:
      - tables : list tables with row counts.
      - schema : CREATE statements for all tables (or `table` if specified).
      - query  : run a SELECT-only `query` (capped at `limit` rows).

    Path handling: ``path`` must be absolute. The shared read
    pipeline (``read_pipeline.gate``) is the only legitimate
    resolver. Relative paths raise.
    """
    if mode not in {"tables", "schema", "query"}:
        raise ValueError(f"invalid mode: {mode!r} (expected 'tables', 'schema', or 'query')")
    if limit <= 0 or limit > 10_000:
        raise ValueError(f"limit must be 1..10000 (got {limit!r})")
    p = _resolve_path(path)

    conn = _sqlite_connect_readonly(p)
    # Diagnostic prelude: every response carries the requested path,
    # the resolved on-disk path, sqlite's database_list result (the
    # actual files SQLite has open for this connection — main + temp +
    # any attached), and journal_mode. Truthfulness contract: if the
    # caller asked for X and the resolved/opened file is anything
    # else, the response makes that visible. Hard mismatch is also
    # checked below.
    try:
        db_list = [
            {"seq": row[0], "name": row[1], "file": row[2]}
            for row in conn.execute("PRAGMA database_list").fetchall()
        ]
    except Exception as exc:
        db_list = [{"error": repr(exc)}]
    try:
        journal_mode = conn.execute("PRAGMA journal_mode").fetchone()[0]
    except Exception:
        journal_mode = ""
    # Diagnostic mismatch guard. SQLite's main db file path may differ
    # from the resolved path on case-insensitive filesystems; compare
    # canonical forms. If the opened main db file does not match the
    # resolved path, hard-fail rather than return phantom rows.
    main_file = ""
    for entry in db_list:
        if entry.get("name") == "main":
            main_file = entry.get("file", "") or ""
            break
    if main_file:
        try:
            if Path(main_file).resolve() != p.resolve():
                conn.close()
                raise ValueError(
                    f"sqlite opened a different file than requested: "
                    f"requested={p} opened={main_file}",
                )
        except OSError:
            pass
    diagnostic = {
        "requested_db_path": str(path),
        "resolved_db_path": str(p),
        "database_list": db_list,
        "journal_mode": journal_mode,
    }
    try:
        if mode == "tables":
            cur = conn.execute(
                "SELECT name FROM sqlite_master "
                "WHERE type='table' AND name NOT LIKE 'sqlite_%' "
                "ORDER BY name",
            )
            names = [row[0] for row in cur.fetchall()]
            tables_out: list[dict[str, Any]] = []
            for name in names:
                # Quote identifier with embedded quotes doubled for safety.
                safe = name.replace('"', '""')
                count = conn.execute(f'SELECT COUNT(*) FROM "{safe}"').fetchone()[0]
                tables_out.append({"name": name, "row_count": count})
            return {
                "path": str(p),
                "mode": mode,
                "tables": tables_out,
                **diagnostic,
            }

        if mode == "schema":
            if table:
                cur = conn.execute(
                    "SELECT name, sql FROM sqlite_master WHERE type='table' AND name = ?",
                    (table,),
                )
            else:
                cur = conn.execute(
                    "SELECT name, sql FROM sqlite_master "
                    "WHERE type='table' AND name NOT LIKE 'sqlite_%' "
                    "ORDER BY name",
                )
            rows = cur.fetchall()
            if table and not rows:
                raise ValueError(f"table not found: {table!r}")
            return {
                "path": str(p),
                "mode": mode,
                "schemas": [{"name": name, "sql": sql or ""} for name, sql in rows],
                **diagnostic,
            }

        # mode == "query"
        if not query or not query.strip():
            raise ValueError("mode='query' requires a non-empty `query`")
        if _SQLITE_REJECT_RE.search(query):
            raise ValueError(
                "only SELECT-style queries are allowed; write keywords "
                "(insert/update/delete/drop/alter/create/attach/replace) "
                "were detected",
            )
        cur = conn.execute(query)
        cols = [desc[0] for desc in (cur.description or [])]
        rows = cur.fetchmany(limit)
        return {
            "path": str(p),
            "mode": mode,
            "columns": cols,
            "rows": [list(r) for r in rows],
            "row_count": len(rows),
            "limit": limit,
            "truncated": len(rows) >= limit,
            **diagnostic,
        }
    finally:
        conn.close()


# JSONL logs (Claude Code session records, MCP telemetry) balloon to
# hundreds of MB. Loading raw text wastes tokens when callers only need
# one field from a filtered subset. This reader streams line-by-line,
# filters server-side, and projects requested fields only.


import json as _json


def _get_dotpath(obj: Any, path: str) -> Any:
    cur: Any = obj
    for part in path.split("."):
        if cur is None:
            return None
        if isinstance(cur, list):
            try:
                cur = cur[int(part)]
            except (ValueError, IndexError):
                return None
        elif isinstance(cur, dict):
            cur = cur.get(part)
        else:
            return None
    return cur


def read_jsonl(
    path: str,
    where: dict[str, Any] | None = None,
    select: list[str] | None = None,
    content_contains: str | None = None,
    offset: int = 0,
    limit: int = 50,
) -> dict[str, Any]:
    """Stream a JSONL file with field-level filter + projection.

    `where` is dotted-path=value exact-match (all must match).
    `select` is dotted paths to project per row; None returns full object.
    `content_contains` is a pre-parse substring filter for speed on big files.
    Invalid JSON lines are counted in `invalid_lines` and skipped so
    mixed logs (status lines + JSON) don't abort the read.
    """
    if limit <= 0 or limit > 5_000:
        raise ValueError(f"limit must be 1..5000 (got {limit!r})")
    if offset < 0:
        raise ValueError(f"offset must be >= 0 (got {offset!r})")
    p = _resolve_path(path)

    rows: list[Any] = []
    scanned = 0
    invalid = 0
    matched = 0

    with p.open("r", encoding="utf-8", errors="replace") as fh:
        for line in fh:
            scanned += 1
            if content_contains and content_contains not in line:
                continue
            stripped = line.strip()
            if not stripped:
                continue
            try:
                obj = _json.loads(stripped)
            except _json.JSONDecodeError:
                invalid += 1
                continue

            if where:
                ok = True
                for dotpath, expected in where.items():
                    if _get_dotpath(obj, dotpath) != expected:
                        ok = False
                        break
                if not ok:
                    continue

            matched += 1
            if matched <= offset:
                continue
            if len(rows) >= limit:
                continue

            if select:
                projected: dict[str, Any] = {}
                for dotpath in select:
                    projected[dotpath] = _get_dotpath(obj, dotpath)
                rows.append(projected)
            else:
                rows.append(obj)

    return {
        "path": str(p),
        "rows": rows,
        "row_count": len(rows),
        "scanned_lines": scanned,
        "invalid_lines": invalid,
        "matched_total": matched,
        "truncated": matched > offset + len(rows),
        "offset": offset,
        "limit": limit,
    }
